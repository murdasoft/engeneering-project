#!/usr/bin/env python3
"""
InspectAI Engineering Platform — Comprehensive Handbook
Generates a large .docx documentation file covering the entire project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "InspectAI-Handbook.docx"

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_paragraph(doc, text, style="Normal"):
    return doc.add_paragraph(text, style=style)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = "Normal"
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    return p

def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")

def set_table_style(table):
    table.style = "Light Grid Accent 1"

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("InspectAI")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x73, 0x77)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Engineering Handbook & Complete Project Documentation")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_paragraph()
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version 1.0 · August 2026")
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("A comprehensive guide for developers, engineers, and operators learning and maintaining the InspectAI structural defect detection platform.")
    run.font.size = Pt(11)

    doc.add_page_break()

    # Table of contents placeholder
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. What InspectAI Does",
        "3. High-Level Architecture",
        "4. Technology Stack",
        "5. Repository Structure",
        "6. The Web Application",
        "7. Backend API Routes",
        "8. Database Schema",
        "9. The ML Service",
        "10. The Telegram Bot",
        "11. Engineering Tool Ecosystem",
        "12. Authentication & Security",
        "13. File Upload & Storage",
        "14. Analysis Pipeline",
        "15. Review Queue & Human-in-the-loop",
        "16. PDF Report Generation",
        "17. Deployment Guide",
        "18. Development Workflow",
        "19. Environment Variables",
        "20. Troubleshooting",
        "21. Glossary",
    ]
    for item in toc_items:
        add_paragraph(doc, item)
    doc.add_page_break()

    # Chapter 1
    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(doc,
        "InspectAI is an AI-powered structural defect detection and engineering analysis platform. "
        "It allows engineers, inspectors, and construction professionals to upload photos of concrete, "
        "brick, masonry, or asphalt surfaces, automatically detect cracks, spalling, corrosion, and other "
        "defects using a YOLOv8 ensemble, and produce engineering reports with annotated images, dimensions, "
        "severity scores, and normative references."
    )
    add_paragraph(doc,
        "This handbook is written as a learning resource. It covers every page, component, function, API route, "
        "service, and machine-learning step in the project. The goal is to enable a new developer or engineer to "
        "understand the entire system from a single document."
    )

    # Chapter 2
    add_heading(doc, "2. What InspectAI Does", level=1)
    add_paragraph(doc, "The platform provides a complete workflow from photo to report:")
    add_numbered(doc, "Upload — A user uploads one or more inspection photos to a project.")
    add_numbered(doc, "Detect — An ensemble of YOLOv8 models analyzes each photo for structural defects.")
    add_numbered(doc, "Review — The user reviews AI-detected findings in a human-in-the-loop queue.")
    add_numbered(doc, "Report — The system generates a professional PDF report with all findings and recommendations.")
    add_paragraph(doc,
        "InspectAI does not replace an engineer. It is a screening tool for preliminary visual assessment. "
        "Every report includes a disclaimer per GOST 31937-2011 stating that the result is not a substitute for "
        "a full instrumental inspection."
    )

    # Chapter 3
    add_heading(doc, "3. High-Level Architecture", level=1)
    add_paragraph(doc, "The project is organized as a monorepo with three main components:")
    add_bullet(doc, "apps/web — Next.js 14 web application and backend API.")
    add_bullet(doc, "apps/ml-service — FastAPI Python service running YOLOv8 inference.")
    add_bullet(doc, "bot/ — Python Telegram bot with AI chat and InspectAI photo analysis.")
    add_bullet(doc, "ecosystem/ — Collection of static HTML engineering calculation tools loaded as iframes.")
    add_paragraph(doc,
        "The web app hosts the primary user interface, stores data in a PostgreSQL database via Prisma, "
        "uploads images to Vercel Blob, and proxies ML requests to the HuggingFace-hosted ML service. "
        "The Telegram bot provides a conversational interface and can analyze photos sent by users."
    )

    # Architecture diagram as text
    add_code_block(doc,
"""User
 |
 +-- Web (Next.js + React) -----> PostgreSQL (Prisma)
 |     |
 |     +---> Vercel Blob (images, PDFs)
 |     |
 |     +---> ML Service (HuggingFace / FastAPI / YOLOv8)
 |
 +-- Telegram Bot (FastAPI webhook)
       |
       +---> ML Service
""")

    # Chapter 4
    add_heading(doc, "4. Technology Stack", level=1)
    add_heading(doc, "4.1 Web Application", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Technology"
    hdr_cells[1].text = "Purpose"
    rows = [
        ("Next.js 14.2.3", "React framework with App Router and API routes"),
        ("React 18.3.1", "UI component library"),
        ("TypeScript 5", "Type-safe development"),
        ("Tailwind CSS 3.4.17", "Utility-first styling"),
        ("NextAuth 4.24.7", "Session-based authentication"),
        ("Prisma 5.22.0", "Database ORM and migrations"),
        ("PostgreSQL", "Relational database"),
        ("@vercel/blob", "Object storage for images and PDFs"),
        ("@react-pdf/renderer", "PDF report generation"),
        ("Recharts", "Dashboard charts"),
        ("Material Symbols", "Icon font"),
    ]
    for tech, purpose in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose

    add_heading(doc, "4.2 ML Service", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Technology"
    hdr_cells[1].text = "Purpose"
    rows = [
        ("FastAPI 0.111.0", "Python API framework"),
        ("Uvicorn", "ASGI server"),
        ("Ultralytics YOLOv8", "Object detection and segmentation"),
        ("OpenCV", "Classical computer vision validation"),
        ("Pillow / PIL", "Image drawing and manipulation"),
        ("NumPy", "Array and image math"),
        ("HuggingFace Hub", "Model download and deployment"),
        ("ReportLab", "PDF report generation"),
    ]
    for tech, purpose in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose

    add_heading(doc, "4.3 Telegram Bot", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Technology"
    hdr_cells[1].text = "Purpose"
    rows = [
        ("FastAPI", "Webhook server"),
        ("python-telegram-bot / custom client", "Telegram API wrapper"),
        ("httpx", "Async HTTP client to ML service"),
        ("Groq / Together AI", "LLM responses for general chat"),
        ("Pydantic Settings", "Typed environment configuration"),
    ]
    for tech, purpose in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose

    # Chapter 5
    add_heading(doc, "5. Repository Structure", level=1)
    add_paragraph(doc, "The repository is located at /Users/aleksandr/engeenering-ml-fasad. The top-level layout is:")
    add_code_block(doc,
"""engeenering-ml-fasad/
├── README.md
├── .env.example
├── main.py                  # Telegram bot FastAPI entry point
├── run_polling.py           # Optional polling runner
├── requirements.txt
├── vercel.json
├── apps/
│   ├── ml-service/
│   │   ├── main.py
│   │   ├── report_generator.py
│   │   ├── requirements.txt
│   │   ├── Dockerfile
│   │   └── models/
│   ├── telegram-bot/        # (currently empty placeholder)
│   └── web/
│       ├── app/             # Next.js App Router pages and API
│       ├── lib/             # Utilities, auth, prisma client
│       ├── prisma/          # Schema and seed
│       ├── public/          # Static assets
│       ├── package.json
│       └── tailwind.config.ts
├── bot/                     # Telegram bot source
│   ├── config.py
│   ├── handlers.py
│   ├── ai_agent.py
│   ├── content.py
│   ├── crm.py
│   ├── inspectai.py
│   ├── lang_detect.py
│   ├── models.py
│   ├── sessions.py
│   ├── stt.py
│   └── telegram_client.py
└── ecosystem/               # Static HTML engineering tools
    ├── crackcalc/
    ├── concretemix/
    ├── loadbear/
    ├── normbase/
    ├── rebardesign/
    └── qa/
""")

    # Chapter 6
    add_heading(doc, "6. The Web Application", level=1)
    add_paragraph(doc,
        "The web application is built with Next.js 14 App Router. It uses server-side rendering where possible, "
        "client components for interactive pages, and Next.js API routes for backend logic. The app is styled "
        "with a custom Tailwind configuration based on a Material Design 3 color system."
    )

    add_heading(doc, "6.1 Root Layout", level=2)
    add_paragraph(doc,
        "File: app/web/app/layout.tsx. The root layout wraps the entire application with the Providers component, "
        "loads the Inter font and Material Symbols icon font from Google Fonts, and sets the base HTML classes. "
        "It uses suppressHydrationWarning to avoid hydration mismatches."
    )
    add_code_block(doc,
"""export const metadata: Metadata = {
  title: "InspectAI — Engineering Hub",
  description: "AI-powered structural defect detection and engineering analysis platform",
};""")

    add_heading(doc, "6.2 Landing Page (/)", level=2)
    add_paragraph(doc,
        "File: app/web/app/page.tsx. The landing page is a marketing-style page explaining the product. "
        "It includes a sticky header, hero section with a simulated analysis card, \"How it works\" steps, "
        "defect type cards, benefits, use cases, FAQ accordion, CTA section, and footer."
    )
    add_paragraph(doc, "Key UI elements and their purpose:")
    add_bullet(doc, "Header — Logo, navigation links (How it works, Defect types, Use cases, FAQ), Demo and Sign in buttons.")
    add_bullet(doc, "Hero — Value proposition, two CTA buttons leading to /login and /demo, animated mock analysis card.")
    add_bullet(doc, "How it works — Four-step explainer: Upload, AI detection, Verification, Report.")
    add_bullet(doc, "Defect types — Cards for Cracks, Spalling, Corrosion, Scaling, Exposed rebar, Efflorescence.")
    add_bullet(doc, "Benefits — Fast, Visual, Reliable, Convenient.")
    add_bullet(doc, "Use cases — Bridge inspection, Building assessment, Road maintenance, Industrial facilities.")
    add_bullet(doc, "FAQ — Accordion with common questions about formats, accuracy, manual editing, and reports.")

    add_heading(doc, "6.3 Authentication — Login (/login)", level=2)
    add_paragraph(doc,
        "File: app/web/app/login/page.tsx. A split-screen page. On the left is a visual panel with the InspectAI "
        "brand and feature list. On the right is the login form. The form uses React state for email and password, "
        "calls signIn from next-auth/react with the credentials provider, and redirects to /dashboard on success."
    )
    add_paragraph(doc, "Key buttons and fields:")
    add_bullet(doc, "Email input — type=email, placeholder 'engineer@inspectai.dev'.")
    add_bullet(doc, "Password input — type=password, placeholder dots.")
    add_bullet(doc, "SIGN IN button — disabled while loading, shows spinner and 'SIGNING IN...' text.")
    add_bullet(doc, "Create one link — routes to /register.")

    add_heading(doc, "6.4 Registration Page (/register)", level=2)
    add_paragraph(doc,
        "File: app/web/app/register/page.tsx. Collects name, email, password, confirm password, phone, company, "
        "and position. On submit it posts to /api/auth/register, then automatically signs the user in with the "
        "new credentials and redirects to the dashboard."
    )

    add_heading(doc, "6.5 Dashboard Layout", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/layout.tsx. Provides the application shell for all dashboard pages. "
        "It includes a desktop sidebar and a mobile drawer, navigation items, a top bar with user avatar, "
        "and redirects unauthenticated users to /login."
    )
    add_paragraph(doc, "Navigation items (navItems array):")
    add_bullet(doc, "Dashboard — /dashboard, icon dashboard")
    add_bullet(doc, "Projects — /dashboard/projects, icon folder")
    add_bullet(doc, "Upload Assets — /dashboard/upload, icon cloud_upload")
    add_bullet(doc, "Analysis Engine — /dashboard/analysis, icon psychology")
    add_bullet(doc, "Review Queue — /dashboard/review, icon fact_check")
    add_bullet(doc, "Final Reports — /dashboard/reports, icon description")
    add_bullet(doc, "Engineering Tools — /dashboard/tools, icon construction")
    add_bullet(doc, "Knowledge Base — /dashboard/knowledge, icon menu_book")
    add_paragraph(doc, "Tool sub-navigation (toolItems array):")
    add_bullet(doc, "CrackCalc — /dashboard/tools/crackcalc")
    add_bullet(doc, "LoadBear — /dashboard/tools/loadbear")
    add_bullet(doc, "ConcreteMix — /dashboard/tools/concretemix")
    add_bullet(doc, "RebarDesign — /dashboard/tools/rebardesign")
    add_bullet(doc, "NormBase — /dashboard/tools/normbase")

    add_heading(doc, "6.6 Dashboard Home (/dashboard)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/page.tsx. The main console after login. It fetches /api/dashboard and shows "
        "KPIs, a severity distribution chart, a 7-day activity timeline, and a recent projects table."
    )
    add_paragraph(doc, "KPI cards:")
    add_bullet(doc, "TOTAL ANALYSES — total number of analyses for the user.")
    add_bullet(doc, "PENDING REVIEWS — findings waiting for human review; highlighted if > 0.")
    add_bullet(doc, "CONFIRMED DEFECTS — findings marked CONFIRMED.")
    add_bullet(doc, "GENERATED REPORTS — reports created for the user's projects.")
    add_paragraph(doc, "Each KPI has a trend indicator (computed from real data) and a progress bar.")

    add_heading(doc, "6.7 Projects List (/dashboard/projects)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/projects/page.tsx. Lists all projects for the logged-in user. "
        "Each project card shows name, site ID, status, photo count, report count, and a GENERATE PDF button. "
        "Clicking a card navigates to the project detail page."
    )
    add_paragraph(doc, "Create project modal fields:")
    add_bullet(doc, "PROJECT NAME — required.")
    add_bullet(doc, "SITE ID — optional identifier for the inspection site.")
    add_bullet(doc, "OBJECT TYPE — dropdown: Bridge, Facade, Pier, Building, Road, Dam, Tunnel.")
    add_bullet(doc, "ADDRESS — optional location.")
    add_bullet(doc, "DESCRIPTION — optional project scope.")

    add_heading(doc, "6.8 Project Detail (/dashboard/projects/[id])", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/projects/[id]/page.tsx. Shows detailed information about a project, "
        "including its inspection photos (assets), completed analyses, and generated reports. "
        "It also contains the analysis context parameters form and the photo gallery."
    )
    add_paragraph(doc, "Analysis parameters (passed to ML API):")
    add_bullet(doc, "PIXEL SCALE (mm/px) — converts pixel dimensions to millimeters.")
    add_bullet(doc, "ENVIRONMENT — atmospheric, indoor, or aggressive.")
    add_bullet(doc, "AGGRESSION — normal, aggressive, or mild.")
    add_bullet(doc, "STRUCTURE TYPE — free text.")
    add_bullet(doc, "CONCRETE GRADE — free text.")
    add_bullet(doc, "REBAR CLASS — free text.")
    add_bullet(doc, "STRUCTURE AGE — free text.")
    add_bullet(doc, "PROTECTIVE LAYER (mm) — concrete cover thickness.")
    add_bullet(doc, "CONFIDENCE THRESHOLD — 0.05 to 0.95 slider.")
    add_paragraph(doc,
        "For each asset, the user can click the image to open AssetModal or click RUN AI ANALYSIS. "
        "There is a GENERATE PROJECT PDF button that creates a report for the whole project."
    )

    add_heading(doc, "6.9 AssetModal", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/projects/[id]/AssetModal.tsx. A modal that opens when a photo is clicked. "
        "It loads the latest COMPLETED analysis and displays the image with bounding boxes or polygon overlays. "
        "It lists findings with class, severity, dimensions, and confidence, allows class/severity filtering, "
        "and supports reviewing each finding with Confirm, Reject, or Edit + reviewer note."
    )
    add_paragraph(doc, "Special actions:")
    add_bullet(doc, "CrackCalc button — opens the crack calculation tool pre-filled with the finding dimensions.")
    add_bullet(doc, "Review icon — opens a textarea and Confirm/Reject/Edit buttons.")

    add_heading(doc, "6.10 Batch Upload (/dashboard/upload)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/upload/page.tsx. Allows the user to upload many photos at once, "
        "select a target project, optionally set analysis parameters, and process each file (upload + analyze) "
        "sequentially. It shows a progress bar and per-file status.")
    add_paragraph(doc, "User flow:")
    add_bullet(doc, "Drag & drop or click to select files (JPG, PNG, WEBP).")
    add_bullet(doc, "Select project from the dropdown.")
    add_bullet(doc, "Expand ANALYSIS PARAMETERS to set pixel scale, environment, aggression, and threshold.")
    add_bullet(doc, "Click UPLOAD & ANALYZE.")
    add_bullet(doc, "Each file is uploaded to /api/assets/upload, then analyzed via /api/analyses/run.")
    add_bullet(doc, "Clicking a result thumbnail opens a detail view with overlays and findings.")

    add_heading(doc, "6.11 Analysis List (/dashboard/analysis)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/analysis/page.tsx. Displays a grid of all analyses for the user. "
        "Each card shows the image, filename, status badge, findings count, confidence, and timestamp. "
        "Clicking a card navigates to /dashboard/analysis/[id].")

    add_heading(doc, "6.12 Analysis Detail (/dashboard/analysis/[id])", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/analysis/[id]/page.tsx. A standalone page for a single analysis. "
        "It shows the image with overlays, severity summary, filters, and a list of findings. "
        "The user can confirm, reject, or edit each finding and add a reviewer note. "
        "It also links to CrackCalc for crack findings.")

    add_heading(doc, "6.13 Review Queue (/dashboard/review)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/review/page.tsx. Human-in-the-loop interface. "
        "Findings are grouped by review status: PENDING, CONFIRMED, REJECTED, EDITED. "
        "Each item shows a thumbnail, class, severity, confidence, and filename. "
        "For PENDING findings the user can Confirm, Reject, or Edit. Editing opens a textarea for a reviewer note.")

    add_heading(doc, "6.14 Reports (/dashboard/reports)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/reports/page.tsx. Lists generated PDF reports. "
        "The user can generate a new report by selecting a project from the modal. "
        "If the report has a stored reportUrl, a DOWNLOAD PDF link is shown; otherwise a REGENERATE button is shown.")

    add_heading(doc, "6.15 Settings (/dashboard/settings)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/settings/page.tsx. User profile and preferences. "
        "It loads the current user from /api/user and allows editing name, phone, company, position, and bio. "
        "It also contains the ML Configuration section with the API endpoint (read-only) and confidence threshold slider, "
        "and a Notifications section with four checkboxes.")
    add_paragraph(doc, "Notification options:")
    add_bullet(doc, "Email notifications for completed analyses")
    add_bullet(doc, "Critical defect alerts")
    add_bullet(doc, "Weekly summary reports")
    add_bullet(doc, "Team activity updates")

    add_heading(doc, "6.16 Engineering Tools (/dashboard/tools and /dashboard/tools/[tool])", level=2)
    add_paragraph(doc,
        "Files: app/dashboard/tools/page.tsx and app/dashboard/tools/[tool]/page.tsx. "
        "The tools index shows cards for CrackCalc, LoadBear, ConcreteMix, RebarDesign, and NormBase. "
        "Each tool page loads an iframe pointing to /tools-assets/<tool> with a 15-second timeout and an error fallback. "
        "If the tool fails to load, the user sees 'Tool unavailable' with a TRY OPEN DIRECTLY button.")

    add_heading(doc, "6.17 Knowledge Base (/dashboard/knowledge)", level=2)
    add_paragraph(doc,
        "File: app/web/app/dashboard/knowledge/page.tsx. A reference guide with three tabs: "
        "Facade Systems (ventilated, wet, structural glazing, curtain walls, precast, stone), "
        "Defect Types (surface/structural cracks, spalling, efflorescence, delamination, corrosion staining, joint failure, biological growth), "
        "and Standards (GOST 31937-2011, EN 1992-1-1, ACI 562-16, ISO 13822).")

    add_heading(doc, "6.18 Public Demo (/demo)", level=2)
    add_paragraph(doc,
        "File: app/web/app/demo/page.tsx. A public page where visitors can upload a single image, run AI analysis, "
        "view bounding boxes/polygons, filter by class and severity, and download a PDF report. "
        "It includes controls for pixel scale and confidence threshold. "
        "The page does not require login and is intended for marketing and quick trials.")

    # Chapter 7
    add_heading(doc, "7. Backend API Routes", level=1)
    add_paragraph(doc,
        "The Next.js app includes API routes under app/api. Each route is a server-side handler that uses "
        "getServerSession for authentication and the Prisma client for database access.")

    add_heading(doc, "7.1 Authentication (/api/auth/[...nextauth])", level=2)
    add_paragraph(doc,
        "File: app/web/app/api/auth/[...nextauth]/route.ts. Implements the NextAuth credentials provider. "
        "It uses bcrypt to compare the supplied password with the stored passwordHash and returns a JWT session.")

    add_heading(doc, "7.2 Registration (/api/auth/register)", level=2)
    add_paragraph(doc,
        "File: app/web/app/api/auth/register/route.ts. Creates a new user with a bcrypt-hashed password. "
        "Returns 409 if the email already exists.")

    add_heading(doc, "7.3 User (/api/user)", level=2)
    add_paragraph(doc,
        "File: app/web/app/api/user/route.ts. GET returns the current user's profile including preferences. "
        "PATCH updates name, phone, company, position, bio, and the JSON preferences object.")

    add_heading(doc, "7.4 Projects (/api/projects and /api/projects/[id])", level=2)
    add_paragraph(doc,
        "GET /api/projects returns all projects for the authenticated user with asset and report counts. "
        "POST creates a new project. /api/projects/[id] GET returns the project with assets and reports; "
        "DELETE removes the project and cascades to assets, analyses, findings, and reports.")

    add_heading(doc, "7.5 Assets (/api/assets/upload)", level=2)
    add_paragraph(doc,
        "File: app/web/app/api/assets/upload/route.ts. Accepts a multipart form with file and projectId, "
        "uploads the file to Vercel Blob, creates an Asset record in the database, and returns the asset.")

    add_heading(doc, "7.6 Analyses (/api/analyses, /api/analyses/[id], /api/analyses/run)", level=2)
    add_paragraph(doc,
        "GET /api/analyses lists the user's analyses. GET /api/analyses/[id] returns a single analysis with asset and findings. "
        "POST /api/analyses/run is the core ML invocation route. It creates an Analysis in PROCESSING state, "
        "fetches the image from Vercel Blob, sends it to the ML /predict/detailed endpoint, "
        "and stores the results as findings. If any step fails the analysis is marked FAILED.")
    add_paragraph(doc, "Parameter whitelist passed to ML:")
    add_code_block(doc,
"""[pixel_scale_mm, environment, aggression, structure_type, concrete_grade,
 rebar_class, structure_age, protective_layer_mm, threshold]""")

    add_heading(doc, "7.7 Findings (/api/findings, /api/findings/[id])", level=2)
    add_paragraph(doc,
        "GET /api/findings?status=PENDING returns findings filtered by review status, scoped to the user's analyses. "
        "PATCH /api/findings/[id] updates reviewStatus and reviewerNote after verifying the finding belongs "
        "to an analysis owned by the current user.")

    add_heading(doc, "7.8 Reports (/api/reports and /api/reports/generate)", level=2)
    add_paragraph(doc,
        "GET /api/reports lists reports for the user's projects. POST /api/reports/generate collects the project, "
        "its assets, analyses, and findings, renders a PDF with @react-pdf/renderer, uploads the PDF to Vercel Blob, "
        "saves the reportUrl in the Report record, and returns the PDF bytes for download.")

    add_heading(doc, "7.9 Dashboard Summary (/api/dashboard)", level=2)
    add_paragraph(doc,
        "File: app/web/app/api/dashboard/route.ts. Aggregates counts: total analyses, pending reviews, confirmed defects, "
        "generated reports, total findings, severity distribution, 7-day activity, and weekly trends compared to the previous week.")

    add_heading(doc, "7.10 ML Proxy (/api/ml/predict and /api/ml/report)", level=2)
    add_paragraph(doc,
        "These routes proxy public/unauthenticated requests to the ML service. They are used by the demo page. "
        "They pass the image file as multipart and forward whitelisted query parameters. /api/ml/report returns a PDF stream.")

    # Chapter 8
    add_heading(doc, "8. Database Schema", level=1)
    add_paragraph(doc,
        "The database is PostgreSQL and is managed by Prisma. The schema is in app/web/prisma/schema.prisma.")
    add_code_block(doc,
"""model User {
  id           String    @id @default(cuid())
  email        String    @unique
  name         String?
  passwordHash String?
  role         Role      @default(ENGINEER)
  phone        String?
  company      String?
  position     String?
  bio          String?
  avatarUrl    String?
  preferences  Json?
  createdAt    DateTime  @default(now())
  updatedAt    DateTime  @updatedAt
  projects     Project[]
  analyses     Analysis[]
  accounts     Account[]
  sessions     Session[]
}""")
    add_paragraph(doc,
        "Key models and relationships:")
    add_bullet(doc, "User — owns projects, analyses, and sessions. Stores profile and preferences JSON.")
    add_bullet(doc, "Project — belongs to a user, has assets and reports.")
    add_bullet(doc, "Asset — belongs to a project, has analyses.")
    add_bullet(doc, "Analysis — belongs to an asset and a user, has findings and stores resultData JSON.")
    add_bullet(doc, "Finding — belongs to an analysis, includes class, confidence, severity, bbox JSON, dimensions, reviewStatus, and reviewerNote.")
    add_bullet(doc, "Report — belongs to a project, has a reportUrl and summary.")
    add_bullet(doc, "Account and Session — NextAuth adapter tables.")

    # Chapter 9
    add_heading(doc, "9. The ML Service", level=1)
    add_paragraph(doc,
        "The ML service is a FastAPI application located in apps/ml-service/main.py. It loads YOLOv8 models from "
        "HuggingFace or local paths, runs ensemble inference, applies classical computer vision validation, "
        "extracts polygons, and generates annotated images and PDF reports.")

    add_heading(doc, "9.1 Environment Configuration", level=2)
    add_paragraph(doc, "The ML service reads the following environment variables:")
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    table.rows[0].cells[0].text = "Variable"
    table.rows[0].cells[1].text = "Description"
    env_rows = [
        ("HF_TOKEN", "HuggingFace authentication token"),
        ("HF_MODEL", "Primary YOLO model repository"),
        ("HF_MODEL_FILE", "Primary model file name (default best.pt)"),
        ("HF_MODEL_FALLBACK", "Fallback model repository"),
        ("HF_MODEL_SECONDARY", "Secondary model repository"),
        ("CONFIDENCE_THRESHOLD", "Default confidence threshold (0.30)"),
        ("MIN_CRACK_AREA", "Minimum area for crack detection (120)"),
        ("MIN_CRACK_MAX_DIM", "Minimum max dimension (30)"),
        ("SMALL_CRACK_AREA_CUTOFF", "NMS permissive cutoff (1500)"),
        ("ML_API_KEY", "API key for protected endpoints"),
        ("ALLOWED_ORIGINS", "CORS allowed origins"),
    ]
    for var, desc in env_rows:
        cells = table.add_row().cells
        cells[0].text = var
        cells[1].text = desc

    add_heading(doc, "9.2 Model Loading", level=2)
    add_paragraph(doc,
        "The get_model() function lazily loads YOLO models into a global cache. It tries local files first, "
        "then downloads from HuggingFace. The supported model keys are primary, fallback, secondary, and general. "
        "The general model is yolov8n.pt and is used to detect non-crack objects that can cause false positives.")

    add_heading(doc, "9.3 Allowed and Rejected Classes", level=2)
    add_paragraph(doc, "The _is_crack_like() function accepts classes such as:")
    add_code_block(doc,
"""crack, cracks, fissure, fracture, split, break, pothole,
damage, defect, hole, spalling, spall, delamination,
rust, rusting, ruststain, corrosion, scaling, efflorescence,
discoloration, stain""")
    add_paragraph(doc, "It rejects labels like background, wall, concrete, surface, normal, good, ok, none, no crack.")
    add_paragraph(doc,
        "The _is_non_crack_object() function filters out pipes, doors, windows, outlets, shadows, stains, joints, "
        "fences, lights, bricks, tiles, panels, and other common objects detected by the general YOLO model.")

    add_heading(doc, "9.4 Computer Vision Validation", level=2)
    add_paragraph(doc,
        "After YOLO inference, each candidate region is validated with classical OpenCV checks:")
    add_bullet(doc, "_edge_density_check — uses Canny edge detection to measure edge pixels inside the box.")
    add_bullet(doc, "_directional_contrast_check — computes Sobel gradients and measures anisotropy; cracks have strong gradient in one direction.")
    add_bullet(doc, "_local_contrast_check — checks for bimodal intensity distribution (dark crack on light background).")
    add_bullet(doc, "_validate_crack_region — combines edge density, anisotropy, and contrast into a single score; uses a lower threshold for small regions.")

    add_heading(doc, "9.5 Ensemble Pipeline", level=2)
    add_paragraph(doc,
        "The run_ensemble() function is the core detection pipeline. It performs the following steps:")
    add_numbered(doc, "Detect non-crack objects using the general YOLO model.")
    add_numbered(doc, "Run the primary YOLO model and filter crack-like classes.")
    add_numbered(doc, "Run the secondary model for additional thin-crack candidates.")
    add_numbered(doc, "Run the fallback model if the first two produce few detections.")
    add_numbered(doc, "Apply classical CV validation to each candidate.")
    add_numbered(doc, "Merge and apply size-aware Non-Maximum Suppression (NMS).")
    add_numbered(doc, "Extract polygon masks when available and build engineering analysis.")
    add_paragraph(doc,
        "The pipeline accepts a threshold parameter to override the default confidence threshold, "
        "allowing users to tune sensitivity from the web UI.")

    add_heading(doc, "9.6 API Endpoints", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_style(table)
    table.rows[0].cells[0].text = "Endpoint"
    table.rows[0].cells[1].text = "Method"
    table.rows[0].cells[2].text = "Purpose"
    ep_rows = [
        ("/predict", "POST", "Run ensemble, return detections and annotated image"),
        ("/predict/detailed", "POST", "Run ensemble, return detailed detections with engineering analysis"),
        ("/report", "POST", "Generate and return a PDF inspection report"),
        ("/health", "GET", "Health check"),
    ]
    for ep, method, purpose in ep_rows:
        cells = table.add_row().cells
        cells[0].text = ep
        cells[1].text = method
        cells[2].text = purpose

    add_heading(doc, "9.7 Engineering Analysis", level=2)
    add_paragraph(doc,
        "Each detection is enriched with engineering metadata. The build_engineering_analysis() function estimates:")
    add_bullet(doc, "estimated_width_mm — using pixel_scale_mm and the bounding box width.")
    add_bullet(doc, "estimated_length_mm — using the major axis of the detection.")
    add_bullet(doc, "estimated_area_cm2 — polygon or box area converted to square centimeters.")
    add_bullet(doc, "ru_name — Russian translation of the defect class.")
    add_bullet(doc, "severity — low, medium, high, or critical based on width and class.")
    add_bullet(doc, "overall_condition — NORMAL, SERVICEABLE, LIMITED, or INADMISSIBLE.")

    add_heading(doc, "9.8 PDF Report Generation", level=2)
    add_paragraph(doc,
        "The report_generator.py module builds a professional PDF report. It includes a cover page, summary, "
        "annotated photos, a findings table, engineering recommendations, normative references, and a disclaimer. "
        "The report uses ReportLab for layout and can be rendered directly by the /report endpoint.")

    # Chapter 10
    add_heading(doc, "10. The Telegram Bot", level=1)
    add_paragraph(doc,
        "The Telegram bot is a FastAPI application in the root main.py. It receives webhook updates from Telegram, "
        "processes text, voice, and photo messages, and can hand off to an LLM for general product questions.")

    add_heading(doc, "10.1 Webhook Entry Point", level=2)
    add_paragraph(doc,
        "The /webhook/telegram endpoint receives Telegram updates. It validates the X-Telegram-Bot-Api-Secret-Token, "
        "extracts the message, and routes it:")
    add_bullet(doc, "/start — sends a greeting and city selection.")
    add_bullet(doc, "Photo — downloads the largest photo, sends it to the ML /predict/detailed endpoint, and sends a summary + PDF.")
    add_bullet(doc, "Voice — downloads the voice file, transcribes it with STT, and routes the text through the normal handler.")
    add_bullet(doc, "Text — routes through handle_message().")

    add_heading(doc, "10.2 Conversation Flow", level=2)
    add_paragraph(doc,
        "The bot uses a session store (in-memory with planned Supabase persistence) and a state machine:")
    add_bullet(doc, "START — initial greeting and language/city selection.")
    add_bullet(doc, "CITY_SELECT — asks the user to choose a city.")
    add_bullet(doc, "MENU — shows the main menu with product options.")
    add_bullet(doc, "FAQ / AI — answers product questions using rules or an LLM.")
    add_bullet(doc, "HANDOFF — pauses the bot and alerts a manager if the user asks for a human.")

    add_heading(doc, "10.3 Photo Analysis in Telegram", level=2)
    add_paragraph(doc,
        "The bot/inspectai.py module implements analyze_photo(). It downloads the image, posts it to ML /predict/detailed, "
        "builds a Russian summary with counts and severities, then posts to ML /report and sends the resulting PDF. "
        "It filters out background/wall/concrete classes client-side as an extra guard.")

    add_heading(doc, "10.4 Configuration", level=2)
    add_paragraph(doc, "The bot/config.py uses Pydantic Settings to load:")
    add_bullet(doc, "Telegram token and webhook secret")
    add_bullet(doc, "AI provider (Groq or Together) and API keys")
    add_bullet(doc, "STT provider and settings")
    add_bullet(doc, "Supabase URL and key")
    add_bullet(doc, "Bitrix24 webhook for leads")
    add_bullet(doc, "Google Sheets integration")
    add_bullet(doc, "PDF presentation file_ids")
    add_bullet(doc, "ML service URL and API key")
    add_bullet(doc, "Feature flags: hybrid_ai, fast_faq, faq_guide_llm")

    add_heading(doc, "10.5 AI Agent", level=2)
    add_paragraph(doc,
        "The bot/ai_agent.py file provides LLM integration. It can detect user intent, answer FAQs, "
        "route product requests, and generate leads in the CRM. When hybrid_ai is enabled, general text messages "
        "are answered by a language model; otherwise the bot uses rule-based responses.")

    add_heading(doc, "10.6 CRM and Lead Creation", level=2)
    add_paragraph(doc,
        "The bot/crm.py module sends lead data to Bitrix24 via a webhook. When a user requests contact, "
        "the bot collects city, product interest, and contact info and creates a lead.")

    # Chapter 11
    add_heading(doc, "11. Engineering Tool Ecosystem", level=1)
    add_paragraph(doc,
        "The ecosystem/ directory contains static HTML engineering tools served under /tools-assets. "
        "The web app loads them as iframes. These tools are self-contained and can also be opened directly.")
    table = doc.add_table(rows=1, cols=3)
    set_table_style(table)
    table.rows[0].cells[0].text = "Tool"
    table.rows[0].cells[1].text = "Path"
    table.rows[0].cells[2].text = "Purpose"
    tool_rows = [
        ("CrackCalc", "ecosystem/crackcalc", "Crack width calculation, growth prediction, multi-standard analysis"),
        ("LoadBear", "ecosystem/loadbear", "Load-bearing capacity of reinforced concrete sections"),
        ("ConcreteMix", "ecosystem/concretemix", "Concrete mix design, cost calculation, granulometry"),
        ("RebarDesign", "ecosystem/rebardesign", "Reinforcement design for concrete sections"),
        ("NormBase", "ecosystem/normbase", "Normative database with GOST, SP, SNiP search"),
        ("QA", "ecosystem/qa", "Quality assurance reference"),
    ]
    for tool, path, purpose in tool_rows:
        cells = table.add_row().cells
        cells[0].text = tool
        cells[1].text = path
        cells[2].text = purpose

    # Chapter 12
    add_heading(doc, "12. Authentication & Security", level=1)
    add_paragraph(doc,
        "Authentication in the web app is handled by NextAuth with a credentials provider. "
        "Users register with email and password; the password is hashed with bcrypt. "
        "The JWT token carries user id and role. API routes use getServerSession to verify the session.")
    add_paragraph(doc, "Security practices:")
    add_bullet(doc, "Secrets are stored in environment variables only.")
    add_bullet(doc, ".env files are gitignored.")
    add_bullet(doc, "ML API key can be passed in the X-API-Key header.")
    add_bullet(doc, "Telegram webhook uses X-Telegram-Bot-Api-Secret-Token for verification.")
    add_bullet(doc, "API routes check asset/project ownership before returning or modifying data.")
    add_bullet(doc, "PATCH /api/findings/[id] verifies the finding belongs to the user's analysis.")

    # Chapter 13
    add_heading(doc, "13. File Upload & Storage", level=1)
    add_paragraph(doc,
        "Images and PDFs are stored in Vercel Blob, a serverless object storage. "
        "When a user uploads a photo, the /api/assets/upload route puts the file in a public blob path "
        "and stores the returned blobUrl in the Asset table. Reports generated by /api/reports/generate are "
        "also uploaded to Vercel Blob and the URL is saved in the Report table.")

    add_heading(doc, "13.1 Upload Flow", level=2)
    add_numbered(doc, "User selects a project and drops files on /dashboard/upload.")
    add_numbered(doc, "For each file, the front-end POSTs multipart to /api/assets/upload.")
    add_numbered(doc, "The API route uses @vercel/blob put() to upload the file.")
    add_numbered(doc, "An Asset row is created with filename, blobUrl, fileSize, mimeType, and projectId.")
    add_numbered(doc, "The asset ID is returned and the front-end triggers /api/analyses/run.")

    # Chapter 14
    add_heading(doc, "14. Analysis Pipeline", level=1)
    add_paragraph(doc,
        "The analysis pipeline moves a photo from upload to a list of findings in the database.")
    add_numbered(doc, "User clicks RUN AI ANALYSIS on an asset or UPLOAD & ANALYZE in batch.")
    add_numbered(doc, "POST /api/analyses/run creates an Analysis row with status PROCESSING.")
    add_numbered(doc, "The API fetches the image bytes from the Vercel Blob URL.")
    add_numbered(doc, "It builds a multipart form and posts to ML /predict/detailed with whitelisted parameters.")
    add_numbered(doc, "The ML service returns detections, engineering data, summary, and an annotated image.")
    add_numbered(doc, "The API creates Finding rows for each defect (class !== 'other').")
    add_numbered(doc, "The Analysis status becomes COMPLETED with confidence and modelVersion, or FAILED on error.")

    # Chapter 15
    add_heading(doc, "15. Review Queue & Human-in-the-loop", level=1)
    add_paragraph(doc,
        "Findings start with reviewStatus PENDING. Engineers verify each finding in the Review Queue, "
        "Project Detail AssetModal, or Analysis Detail page. They can:")
    add_bullet(doc, "CONFIRM — accept the AI detection as correct.")
    add_bullet(doc, "REJECT — mark the detection as false positive.")
    add_bullet(doc, "EDIT — mark the detection as modified and optionally add a reviewer note.")
    add_paragraph(doc,
        "Confirmed findings count toward project statistics and appear in PDF reports. "
        "Rejected findings are filtered out. Edited findings are kept with a note for traceability.")

    # Chapter 16
    add_heading(doc, "16. PDF Report Generation", level=1)
    add_paragraph(doc,
        "There are two PDF generation paths: the web app report (React-PDF) and the ML service report (ReportLab). "
        "The web app path is used for project-level reports in the dashboard.")

    add_heading(doc, "16.1 Web Report", level=2)
    add_paragraph(doc,
        "/api/reports/generate collects all assets, analyses, and findings for a project, renders a PDF with "
        "@react-pdf/renderer using lib/pdf-report.tsx, uploads the PDF to Vercel Blob, and returns the bytes "
        "for immediate download. The stored reportUrl allows re-downloading later.")

    add_heading(doc, "16.2 ML Report", level=2)
    add_paragraph(doc,
        "The ML /report endpoint uses report_generator.py. It accepts a single image and project metadata, "
        "runs inference, and produces a Russian-language report with an annotated cover image, summary, "
        "findings list, and engineering recommendations. This is used by the demo page and Telegram bot.")

    # Chapter 17
    add_heading(doc, "17. Deployment Guide", level=1)
    add_heading(doc, "17.1 Web App on Vercel", level=2)
    add_numbered(doc, "Install dependencies: cd apps/web && npm install")
    add_numbered(doc, "Set environment variables in Vercel: DATABASE_URL, DIRECT_URL, NEXTAUTH_SECRET, ML_API_URL, ML_API_KEY.")
    add_numbered(doc, "Run Prisma generate and db push: npx prisma generate && npx prisma db push")
    add_numbered(doc, "Deploy: vercel --prod")

    add_heading(doc, "17.2 ML Service on HuggingFace Spaces", level=2)
    add_numbered(doc, "Create a Docker Space on HuggingFace.")
    add_numbered(doc, "Upload apps/ml-service contents including Dockerfile, main.py, report_generator.py, requirements.txt.")
    add_numbered(doc, "Set HF_TOKEN, ML_API_KEY, and other variables in Space settings.")
    add_numbered(doc, "The Space builds and serves the FastAPI app.")

    add_heading(doc, "17.3 Telegram Bot on Vercel", level=2)
    add_numbered(doc, "Install Python dependencies from requirements.txt.")
    add_numbered(doc, "Set .env variables for Telegram token, webhook secret, AI keys, and ML service URL.")
    add_numbered(doc, "Deploy with vercel deploy --prod.")
    add_numbered(doc, "Register the webhook: curl https://your-domain/setup")

    # Chapter 18
    add_heading(doc, "18. Development Workflow", level=1)
    add_paragraph(doc, "Recommended workflow for contributors:")
    add_numbered(doc, "Clone the repository.")
    add_numbered(doc, "Copy .env.example to .env and fill in real values.")
    add_numbered(doc, "Install web dependencies and run prisma generate.")
    add_numbered(doc, "Start the Next.js dev server: npm run dev")
    add_numbered(doc, "Run the ML service locally with uvicorn main:app --reload --port 8000")
    add_numbered(doc, "For the bot, run uvicorn main:app --reload --port 8000 and use ngrok for webhook testing.")

    add_heading(doc, "18.1 Adding a New Dashboard Page", level=2)
    add_bullet(doc, "Create page.tsx under app/dashboard/<path>.")
    add_bullet(doc, "Add the route and icon to the navItems array in app/dashboard/layout.tsx if it belongs in the main menu.")
    add_bullet(doc, "Create the API route under app/api/<path>/route.ts if server data is needed.")
    add_bullet(doc, "Update the Prisma schema if new data models are required, then run db push.")

    # Chapter 19
    add_heading(doc, "19. Environment Variables", level=1)
    add_paragraph(doc, "A complete list of key environment variables for each service:")
    table = doc.add_table(rows=1, cols=3)
    set_table_style(table)
    table.rows[0].cells[0].text = "Variable"
    table.rows[0].cells[1].text = "Service"
    table.rows[0].cells[2].text = "Purpose"
    var_rows = [
        ("DATABASE_URL", "Web", "PostgreSQL connection string"),
        ("DIRECT_URL", "Web", "Direct PostgreSQL connection for migrations"),
        ("NEXTAUTH_SECRET", "Web", "NextAuth JWT encryption"),
        ("ML_API_URL", "Web/Bot", "URL of the ML service"),
        ("ML_API_KEY", "Web/Bot/ML", "API key for ML service"),
        ("HF_TOKEN", "ML", "HuggingFace token"),
        ("HF_MODEL", "ML", "Primary model repo"),
        ("telegram_bot_token", "Bot", "Telegram Bot API token"),
        ("telegram_webhook_secret", "Bot", "Webhook verification"),
        ("webhook_base_url", "Bot", "Public URL for webhook"),
        ("groq_api_key", "Bot", "Groq LLM key"),
        ("together_api_key", "Bot", "Together AI LLM key"),
        ("supabase_url", "Bot", "Supabase URL"),
        ("supabase_service_role_key", "Bot", "Supabase service key"),
        ("bitrix24_webhook_url", "Bot", "Bitrix24 lead webhook"),
    ]
    for var, service, purpose in var_rows:
        cells = table.add_row().cells
        cells[0].text = var
        cells[1].text = service
        cells[2].text = purpose

    # Chapter 20
    add_heading(doc, "20. Troubleshooting", level=1)
    add_paragraph(doc, "Common issues and how to resolve them:")
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    table.rows[0].cells[0].text = "Issue"
    table.rows[0].cells[1].text = "Solution"
    issue_rows = [
        ("Next.js hydration errors", "Move hooks before conditional returns and extract helper components outside the main component."),
        ("Prisma client not found", "Run npx prisma generate after schema changes."),
        ("ML service returns no detections", "Lower the confidence threshold or increase pixel_scale_mm."),
        ("PDF download link missing", "Verify /api/reports/generate uploaded the PDF to Vercel Blob and saved reportUrl."),
        ("Telegram webhook not receiving updates", "Call /setup endpoint and check webhook_base_url and secret token."),
        ("CORS errors from ML", "Set ALLOWED_ORIGINS to the web domain."),
        ("Database migration fails", "Check DATABASE_URL and DIRECT_URL; run db push when the database is reachable."),
    ]
    for issue, solution in issue_rows:
        cells = table.add_row().cells
        cells[0].text = issue
        cells[1].text = solution

    # Chapter 21
    add_heading(doc, "21. Glossary", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    gloss_rows = [
        ("YOLO", "You Only Look Once — a family of object detection models."),
        ("mAP50", "Mean Average Precision at 50% IoU threshold."),
        ("NMS", "Non-Maximum Suppression — removes duplicate overlapping boxes."),
        ("Bbox", "Bounding box — a rectangle around a detected object."),
        ("Polygon", "A multi-point shape outlining a detected defect."),
        ("Pixel scale", "Real-world millimeters per image pixel."),
        ("Confidence threshold", "Minimum detection score to keep a prediction."),
        ("CORS", "Cross-Origin Resource Sharing."),
        ("Prisma", "Type-safe database ORM."),
        ("NextAuth", "Authentication library for Next.js."),
        ("Vercel Blob", "Serverless object storage."),
        ("HuggingFace Spaces", "Cloud platform for ML demos."),
    ]
    for term, definition in gloss_rows:
        cells = table.add_row().cells
        cells[0].text = term
        cells[1].text = definition

    doc.add_page_break()
    add_heading(doc, "Appendix A: File Index", level=1)
    add_paragraph(doc, "This appendix provides a quick reference of the most important files in the project.")
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table)
    table.rows[0].cells[0].text = "File"
    table.rows[0].cells[1].text = "Description"
    file_rows = [
        ("apps/web/app/layout.tsx", "Root layout with fonts and providers"),
        ("apps/web/app/page.tsx", "Marketing landing page"),
        ("apps/web/app/dashboard/layout.tsx", "Dashboard sidebar and shell"),
        ("apps/web/app/dashboard/page.tsx", "Dashboard KPI and charts"),
        ("apps/web/app/dashboard/upload/page.tsx", "Batch upload and analysis"),
        ("apps/web/app/dashboard/projects/[id]/page.tsx", "Project detail page"),
        ("apps/web/app/api/analyses/run/route.ts", "Core ML analysis API"),
        ("apps/web/app/api/reports/generate/route.ts", "PDF report API"),
        ("apps/web/lib/auth.ts", "NextAuth credentials configuration"),
        ("apps/web/prisma/schema.prisma", "Database schema"),
        ("apps/ml-service/main.py", "FastAPI ML service"),
        ("apps/ml-service/report_generator.py", "PDF report builder"),
        ("bot/inspectai.py", "Telegram photo analysis"),
        ("bot/handlers.py", "Telegram conversation orchestrator"),
        ("main.py", "Telegram bot FastAPI entry"),
    ]
    for f, d in file_rows:
        cells = table.add_row().cells
        cells[0].text = f
        cells[1].text = d

    # Expanded deep-dive chapters
    add_expanded_chapters(doc)

    # Extra detailed content
    add_extra_content(doc)

    # Source code reference
    add_code_reference(doc)

    # Final save
    doc.save(OUT_PATH)
    print(f"Documentation saved to {OUT_PATH}")


def add_expanded_chapters(doc):
    """Append deep-dive chapters with code walkthroughs to reach 100+ pages."""

    add_heading(doc, "Part II: Deep Dives", level=1)
    add_paragraph(doc,
        "The following chapters provide extensive, line-by-line and concept-by-concept explanations of the most "
        "important code paths, user interface flows, and service interactions. The goal is to give a student or "
        "new contributor enough context to understand, modify, and extend InspectAI without guessing."
    )

    # Web app deep dive
    add_heading(doc, "22. Web Application in Detail", level=1)
    add_paragraph(doc,
        "The InspectAI web frontend is a Next.js 14 application using the App Router. Every route is defined by a "
        "page.tsx file. Client components use 'use client' at the top. Server components are used for layout when "
        "possible, but the dashboard is interactive and therefore a client component."
    )

    add_heading(doc, "22.1 Landing Page Anatomy", level=2)
    add_paragraph(doc,
        "The landing page (app/web/app/page.tsx) is intentionally self-contained. It uses inline styles rather than "
        "Tailwind to ensure the page renders correctly even if the rest of the design system has issues. It is a "
        "single React functional component HomePage with one useState hook for the FAQ accordion.")
    add_paragraph(doc, "The openFaq state holds the index of the currently open FAQ item. If openFaq === i, the answer "
        "paragraph is rendered with a fade-in animation. Clicking the same question again collapses it by setting "
        "openFaq to null. This is a simple accordion pattern.")
    add_paragraph(doc,
        "The page uses several CSS keyframe animations injected with a style tag: fadeIn, slideUp, float, and pulse. "
        "Elements receive class names like anim-1, anim-2, anim-3, anim-4 for staggered entrance, and card-hover for "
        "hover lift effects. The nav-link class adds an animated underline on hover. The btn-hover class lifts buttons "
        "and adds a shadow.")
    add_paragraph(doc,
        "The hero section uses a grid with two columns on desktop. The left side has the headline, description, and CTA "
        "buttons. The right side shows a mock 'Analysis Result' card with three colored bounding boxes representing a crack, "
        "corrosion, and spalling detection. The boxes are absolutely positioned inside a gray background. A decorative "
        "SVG path shows a wavy crack.")
    add_paragraph(doc,
        "FAQ data is stored as a static array of objects with q and a properties. This makes it easy to translate or extend. "
        "The footer is a three-column flex with the logo, product tagline, and copyright.")

    add_heading(doc, "22.2 Login and Registration Flow", level=2)
    add_paragraph(doc,
        "Both /login and /register use a split-screen layout. On desktop, the left side shows a branded visual panel. "
        "On mobile, the visual panel is hidden and the logo is shown above the form. This is achieved with Tailwind's "
        "hidden lg:flex and lg:hidden classes.")
    add_paragraph(doc,
        "The login form uses controlled inputs. On submit it calls signIn('credentials', { email, password, redirect: false }). "
        "If res.error is present, the error message is shown. Otherwise the Next.js router pushes /dashboard. The form "
        "disables the submit button while loading and shows a spinning Material Symbols icon.")
    add_paragraph(doc,
        "The registration form has more fields: name, email, password, confirmPassword, phone, company, position. It first "
        "validates that the two password inputs match. It then posts the data to /api/auth/register. If the response is "
        "successful, it signs the user in with the new credentials. This provides a seamless one-step registration to dashboard "
        "experience.")
    add_paragraph(doc,
        "Error messages are rendered in a rounded alert box with bg-error-container and text-on-error-container colors. "
        "Input styling uses the design system's surface and outline colors with focus:outline-none and focus:border-primary.")

    add_heading(doc, "22.3 Dashboard Layout Components", level=2)
    add_paragraph(doc,
        "The dashboard layout (app/web/app/dashboard/layout.tsx) is a client component. It uses useSession from next-auth/react "
        "to get the user's authentication status. If the status is loading or unauthenticated, it shows a spinner. If "
        "unauthenticated after mount, it redirects to /login.")
    add_paragraph(doc,
        "NavLink is a helper component. It receives an item, an onClick handler, an optional isTool flag, and the current "
        "pathname. It computes active based on whether the current pathname matches or starts with the item href. Tool items "
        "use startsWith because they are nested under /dashboard/tools. For active links it applies bg-primary-container, "
        "text-on-primary-container, and a left border. For inactive links it uses text-on-surface-variant and hover background.")
    add_paragraph(doc,
        "SidebarContent renders the main navigation and the engineering tools section. It also renders a bottom section with "
        "a NEW INSPECTION button and Settings / Sign Out links. Sign Out calls signOut({ callbackUrl: '/login' }) from next-auth/react.")
    add_paragraph(doc,
        "The mobile sidebar is shown when mobileOpen is true. It includes an overlay with bg-black/50 and a fixed aside "
        "panel. Clicking the overlay or the close button sets mobileOpen to false. The pathname useEffect closes the mobile "
        "drawer automatically on navigation.")
    add_paragraph(doc,
        "The main content area has a sticky top bar with a mobile hamburger menu, the 'ENGINEERING CONSOLE' label, a 'LIVE' "
        "indicator, and a circular avatar with the user's first initial. The children prop renders the current page.")

    add_heading(doc, "22.4 Dashboard Home Deep Dive", level=2)
    add_paragraph(doc,
        "The dashboard home (app/web/app/dashboard/page.tsx) is a client component. On mount it fetches /api/dashboard and "
        "stores the result in data. It shows a loading spinner until the data arrives.")
    add_paragraph(doc,
        "The KPI grid uses a 1-column layout on mobile and 4 columns on desktop (grid-cols-1 md:grid-cols-4). Each card has "
        "a label, a large value, a trend indicator, and a progress bar. The PENDING REVIEWS card has a highlight variant with "
        "a 'HIGH ATTENTION' badge when pendingReviews > 0. The color and barWidth are hardcoded per card, but the trend value "
        "and the numeric value come from the API.")
    add_paragraph(doc,
        "The severity distribution section uses a conic-gradient background on a circular div to show the proportion of each "
        "severity. It computes percentages by dividing each severity count by the total. A center overlay shows the total "
        "findings count. Below the circle, a list shows the count and percentage for each severity with a colored dot.")
    add_paragraph(doc,
        "The activity timeline is a bar chart built with plain divs. It maps the 7 entries from data.activity to bars with "
        "heights proportional to the maximum count. The bar for the maximum value uses bg-primary; others use bg-primary/20. "
        "Day labels are derived from the date's UTC day of week.")
    add_paragraph(doc,
        "Recent projects are shown in a table with columns for Project Name, Object Type, Photos, Last Activity, and Status. "
        "Status is rendered as a colored badge. If there are no projects, a row shows a link to create the first project.")
    add_paragraph(doc,
        "A floating action button (FAB) is fixed to the bottom right and links to /dashboard/upload. It uses a primary-container "
        "background and an add_box icon.")

    add_heading(doc, "22.5 Projects List and Creation", level=2)
    add_paragraph(doc,
        "The projects page (app/web/app/dashboard/projects/page.tsx) has a Projects interface and uses useState for the list, "
        "loading state, create modal visibility, generating state, and new project form values. The form values object has "
        "name, siteId, objectType, address, and description.")
    add_paragraph(doc,
        "loadProjects fetches from /api/projects. Each project is rendered as a card inside a grid. The card is also a Link to "
        "/dashboard/projects/[id]. Inside the card, the GENERATE PDF button is an actual button, not a link. It calls e.preventDefault() "
        "and e.stopPropagation() so that clicking it does not navigate to the project detail. The generateReport function posts "
        "to /api/reports/generate, downloads the PDF blob, and triggers a file download in the browser.")
    add_paragraph(doc,
        "The create project modal is a fixed overlay. It stops click propagation so the modal itself does not close when clicked. "
        "It validates the name on submit and posts to /api/projects. On success it clears the form and reloads the list.")
    add_paragraph(doc,
        "Status colors map ACTIVE to primary, COMPLETED to secondary, CRITICAL to error, and ARCHIVED to outline. The card footer "
        "shows object type and creation date.")

    add_heading(doc, "22.6 Project Detail Page", level=2)
    add_paragraph(doc,
        "The project detail page (app/web/app/dashboard/projects/[id]/page.tsx) is one of the most complex client pages. It uses "
        "useParams to read the projectId, fetches the project from /api/projects/[id], and manages analyzing and generating states.")
    add_paragraph(doc,
        "The analysisParams state is a record of strings and numbers with defaults for pixel scale, environment, aggression, "
        "structure type, concrete grade, rebar class, structure age, protective layer, and threshold. These are rendered in a "
        "right-side panel. Changes are stored in React state and passed to /api/analyses/run when the user clicks RUN AI ANALYSIS.")
    add_paragraph(doc,
        "The asset list is rendered as a grid of cards. Each card shows the image, filename, upload date, and existing analyses. "
        "If an analysis is completed, it shows findings count and confidence. If no analysis exists, it shows a RUN AI ANALYSIS "
        "button. Clicking the image opens AssetModal with the asset.")
    add_paragraph(doc,
        "runAnalysis posts { assetId, projectId, params: analysisParams } to /api/analyses/run. If the response is not OK, it "
        "shows an alert. On success it reloads the project. generateReport calls /api/reports/generate with the projectId, "
        "downloads the PDF, and triggers a download.")
    add_paragraph(doc,
        "The page also has a Reports section listing existing project reports and a Summary panel showing total findings by severity.")

    add_heading(doc, "22.7 AssetModal Component", level=2)
    add_paragraph(doc,
        "AssetModal (app/web/app/dashboard/projects/[id]/AssetModal.tsx) is a large client component. It receives an asset and "
        "a callback onClose. It loads the latest COMPLETED analysis for the asset, computes an overlay scale based on the image "
        "natural size and the container width, and renders interactive bounding boxes or polygons.")
    add_paragraph(doc,
        "The component uses filters for class and severity. Filtered findings are still rendered as overlays but with reduced "
        "opacity (dim) if they do not match. Clicking an overlay or a list item sets the activeFinding, which scrolls the list "
        "and highlights the card.")
    add_paragraph(doc,
        "Each finding card shows the class name, severity badge, confidence, dimensions in millimeters, and engineering notes. "
        "There is a 'CrackCalc' button for findings that have width and height; it opens the tool overlay with search params "
        "prefilled. There is also a review icon that opens a review panel with a textarea and Confirm / Reject / Edit / Cancel "
        "buttons. The PATCH request updates reviewStatus and reviewerNote.")
    add_paragraph(doc,
        "The image container is wrapped with a ref. On image load it computes the display scale. Bounding boxes are positioned "
        "absolutely using left/top/width/height percentages. Polygon overlays use inline SVG paths if the finding has a polygon "
        "array. This allows irregular defect shapes to be highlighted.")

    add_heading(doc, "22.8 Batch Upload Page", level=2)
    add_paragraph(doc,
        "The upload page (app/web/app/dashboard/upload/page.tsx) is the main ingestion point. It has a projects dropdown, a "
        "drag-and-drop zone, a file list, analysis parameters, and a result detail view.")
    add_paragraph(doc,
        "handleFiles receives a FileList and creates UploadedFile objects for each file. Each object has a file, name, size, "
        "status, and a previewUrl generated with URL.createObjectURL. The file list is shown with thumbnails and status badges.")
    add_paragraph(doc,
        "uploadAll is an async loop over the files. For each file it sets the status to 'uploading' with the current progress, "
        "posts a FormData to /api/assets/upload, then posts the analysis request to /api/analyses/run. If any step fails, the "
        "status becomes 'error' with an error message. On success the status is 'done' and the analysis summary is stored.")
    add_paragraph(doc,
        "The detail view is shown when the user clicks a completed result. It fetches the analysis detail (using the analysisId "
        "from the result) and displays the image with overlays. The activeTool state is used to open ToolOverlay with CrackCalc "
        "for a selected finding.")
    add_paragraph(doc,
        "The analysis parameters section is collapsible. It includes number inputs, selects, and a threshold slider. These are "
        "all controlled components tied to analysisParams state and are sent with the analysis request.")

    add_heading(doc, "22.9 Review Queue Page", level=2)
    add_paragraph(doc,
        "The review page (app/web/app/dashboard/review/page.tsx) fetches findings from /api/findings?status=filter. The filter "
        "state can be PENDING, CONFIRMED, REJECTED, or EDITED. Clicking a filter button reloads the list.")
    add_paragraph(doc,
        "For PENDING findings, the user sees Confirm, Reject, and Edit buttons. Clicking Edit sets editingId and copies the "
        "existing reviewerNote (if any) into editNote. A textarea appears. The user can SAVE EDIT to set the status to EDITED "
        "with the note, or CANCEL to close the editor. For non-pending filters, if a finding has a reviewerNote, it is displayed "
        "below the filename.")
    add_paragraph(doc,
        "The severityColors map uses Tailwind classes for LOW, MEDIUM, HIGH, and CRITICAL. The list uses a card layout with "
        "a thumbnail, details, and action buttons. After any action, the finding is removed from the local list so the user sees "
        "progress immediately.")

    add_heading(doc, "22.10 Reports Page", level=2)
    add_paragraph(doc,
        "The reports page (app/web/app/dashboard/reports/page.tsx) loads reports and projects in parallel. It has a generate modal "
        "that lists projects and allows the user to generate a PDF for any of them. The generateReport function posts to "
        "/api/reports/generate, downloads the PDF, and then refreshes the reports list.")
    add_paragraph(doc,
        "Each report card shows a PDF icon, title, summary, project link, and a download action. If reportUrl is present, a link "
        "opens the Vercel Blob URL in a new tab. If it is missing (legacy reports), a REGENERATE button calls generateReport for "
        "that report's project.")

    add_heading(doc, "22.11 Settings Page", level=2)
    add_paragraph(doc,
        "The settings page (app/web/app/dashboard/settings/page.tsx) loads the user profile from /api/user and initializes local "
        "state for each editable field. The ML API endpoint is shown disabled. The confidence threshold slider updates the threshold "
        "state and displays the current value. The notification checkboxes update the notifications state object.")
    add_paragraph(doc,
        "saveProfile sends a PATCH to /api/user with the profile fields and a preferences object containing threshold and "
        "notifications. On success it sets saved to true for three seconds to show the SAVED indicator.")

    add_heading(doc, "22.12 Tools and Knowledge Base", level=2)
    add_paragraph(doc,
        "The tools index (app/dashboard/tools/page.tsx) is a static grid of cards. Each card links to /dashboard/tools/[tool]. "
        "The tool page uses an iframe and a 15-second timeout. If the iframe fails to load, a fallback UI is shown. This handles "
        "tools that are not deployed or temporarily unavailable.")
    add_paragraph(doc,
        "The knowledge base (app/dashboard/knowledge/page.tsx) uses a tab state for facades, defects, and standards. It uses "
        "static reference data. This data could be moved to a CMS or database in the future, but for the MVP it is hardcoded in the "
        "component.")

    add_heading(doc, "22.13 Public Demo Page", level=2)
    add_paragraph(doc,
        "The demo page (app/web/app/demo/page.tsx) does not require authentication. It allows a single image upload, shows a preview, "
        "runs analysis through /api/ml/predict, and lets the user download a PDF through /api/ml/report. It has controls for pixel "
        "scale and threshold, which are passed as query parameters to the API routes.")
    add_paragraph(doc,
        "The result view shows the annotated image with bounding boxes. It has a filter for all, crack, and other detections and a "
        "toggle to show or hide boxes. The detection cards list class, confidence, severity, and dimensions. The PDF button "
        "triggers downloadPdf, which posts the same image to /api/ml/report and downloads the returned PDF.")

    # API deep dive
    add_heading(doc, "23. Backend API in Detail", level=1)
    add_paragraph(doc,
        "The API is built with Next.js Route Handlers. Each file under app/api/<path>/route.ts exports one or more HTTP method "
        "handlers. All protected handlers start with getServerSession(authOptions) and return 401 if the user is not authenticated.")

    add_heading(doc, "23.1 Authentication Route", level=2)
    add_paragraph(doc,
        "The NextAuth route is a catch-all handler at app/api/auth/[...nextauth]/route.ts. The auth configuration is in lib/auth.ts. "
        "It uses CredentialsProvider with email and password fields. The authorize callback finds the user by email, compares the "
        "bcrypt hash, and returns the user object. The JWT and session callbacks attach the user id and role to the session so that "
        "API routes can identify the user.")

    add_heading(doc, "23.2 Registration Route", level=2)
    add_paragraph(doc,
        "The register route (app/api/auth/register/route.ts) checks for required email and password, hashes the password with "
        "bcrypt, and creates a user. If the email already exists it returns a 409 Conflict. On success it returns the user without "
        "the password hash.")

    add_heading(doc, "23.3 User Route", level=2)
    add_paragraph(doc,
        "The user route (app/api/user/route.ts) GET returns the current user's profile including the JSON preferences field. The "
        "PATCH route accepts name, phone, company, position, bio, and preferences. It uses a spread pattern to only update fields "
        "that are present in the request body. This prevents the request from overwriting optional fields with undefined.")

    add_heading(doc, "23.4 Projects Routes", level=2)
    add_paragraph(doc,
        "The projects list route (app/api/projects/route.ts) returns projects scoped to the user. The project detail route "
        "(app/api/projects/[id]/route.ts) GET returns a project with its assets and reports. The POST route for creating projects "
        "uses prisma.user.upsert to ensure a User row exists even if the session was created from a non-standard source.")

    add_heading(doc, "23.5 Asset Upload Route", level=2)
    add_paragraph(doc,
        "The upload route (app/api/assets/upload/route.ts) parses a multipart form with formidable or a similar parser. It reads the "
        "file and projectId, calls put() from @vercel/blob with the file buffer and a content type, and creates an Asset row with "
        "the returned blobUrl. The route returns the asset object.")

    add_heading(doc, "23.6 Analysis Run Route", level=2)
    add_paragraph(doc,
        "The analysis run route (app/api/analyses/run/route.ts) is the most important route. It performs the following steps:")
    add_numbered(doc, "Authenticate the user with getServerSession.")
    add_numbered(doc, "Read assetId, projectId, and params from the JSON body.")
    add_numbered(doc, "Verify that the asset exists and belongs to a project owned by the user.")
    add_numbered(doc, "Create an Analysis row with status PROCESSING.")
    add_numbered(doc, "Fetch the image bytes from the Vercel Blob URL.")
    add_numbered(doc, "Build a FormData with the image and append whitelisted query parameters.")
    add_numbered(doc, "Post to the ML /predict/detailed endpoint with the API key header if configured.")
    add_numbered(doc, "Parse the JSON response and create Finding rows for each defect detection.")
    add_numbered(doc, "Update the Analysis to COMPLETED with confidence, modelVersion, and resultData, or to FAILED on error.")
    add_numbered(doc, "Return the analysisId, summary, and findingsCount.")
    add_paragraph(doc,
        "The resultData is stored without the annotated_image field because base64 images are large and the blob URL is used "
        "for display. Findings use the severity value converted to uppercase to match the Prisma enum.")

    add_heading(doc, "23.7 Findings Routes", level=2)
    add_paragraph(doc,
        "The findings list route (app/api/findings/route.ts) accepts a status query parameter and returns findings where the "
        "associated analysis's userId matches the session user. It includes the analysis and asset relations.")
    add_paragraph(doc,
        "The finding update route (app/api/findings/[id]/route.ts) PATCH first finds the finding by id with a where clause that "
        "requires analysis.userId to equal the current user. If no finding is found, it returns 404. This is the ownership check. "
        "If found, it updates reviewStatus and reviewerNote.")

    add_heading(doc, "23.8 Report Generation Route", level=2)
    add_paragraph(doc,
        "The report route (app/api/reports/generate/route.ts) loads the project with assets, analyses, and findings. It computes "
        "statistics: total assets, total analyses, total findings, and counts by severity. It then creates a report data object "
        "with project info, user info, and the findings.")
    add_paragraph(doc,
        "The PDF is rendered with createReportDocument from lib/pdf-report.tsx and renderToBuffer from @react-pdf/renderer. The "
        "resulting buffer is uploaded to Vercel Blob with a public path. The blob URL is saved as reportUrl on the Report row. "
        "Finally the PDF bytes are returned with Content-Type application/pdf and a Content-Disposition attachment header.")

    add_heading(doc, "23.9 Dashboard Route", level=2)
    add_paragraph(doc,
        "The dashboard route (app/api/dashboard/route.ts) uses Promise.all to query several counts at once. It computes the "
        "severity distribution by grouping findings by severity. It builds a 7-day activity array by iterating over analysis "
        "creation dates and matching them to the last 7 days. It also computes week-over-week trend percentages for analyses "
        "and findings.")

    add_heading(doc, "23.10 ML Proxy Routes", level=2)
    add_paragraph(doc,
        "The predict and report proxy routes (app/api/ml/predict/route.ts and app/api/ml/report/route.ts) are intentionally "
        "unauthenticated. They are used by the public demo page. They forward whitelisted query parameters and the uploaded "
        "file to the ML service and stream back the response.")

    # ML deep dive
    add_heading(doc, "24. ML Service in Detail", level=1)
    add_paragraph(doc,
        "The ML service is a single FastAPI file (apps/ml-service/main.py) with helper functions and a report generator module. "
        "It is designed to be deployed as a HuggingFace Space or as a standalone Docker container.")

    add_heading(doc, "24.1 Service Setup", level=2)
    add_paragraph(doc,
        "The FastAPI app is created with a title, description, and version. CORS middleware is added with origins from "
        "ALLOWED_ORIGINS. The service optionally logs into HuggingFace if HF_TOKEN is set. Models are loaded lazily and cached "
        "in the _models global dictionary.")

    add_heading(doc, "24.2 Model Download and Loading", level=2)
    add_paragraph(doc,
        "get_model() first checks if the model is already in the _models cache. If not, it imports YOLO and torch and "
        "temporarily overrides torch.load to set weights_only=False for compatibility. For the primary model it checks local "
        "paths 'best.pt' and 'models/best.pt', then falls back to hf_hub_download, and finally to yolov8n.pt if nothing is "
        "available. For fallback and secondary models it tries to download from HuggingFace and sets None if unavailable. The "
        "general model is always yolov8n.pt.")

    add_heading(doc, "24.3 Class Filters", level=2)
    add_paragraph(doc,
        "_is_crack_like() accepts a class name if it appears in a known list of defect words or contains one of them as a "
        "substring, unless it contains a rejected word like background, wall, concrete, or normal. This allows the model to "
        "return generic labels like 'crack_concrete' while rejecting 'concrete_normal'.")
    add_paragraph(doc,
        "_is_non_crack_object() is used on the general model's outputs. It checks if the class is in a known non-crack set "
        "or if any non-crack word appears in the class name. These objects are stored so that later crack detections overlapping "
        "with them can be removed.")

    add_heading(doc, "24.4 Computer Vision Validation Functions", level=2)
    add_paragraph(doc,
        "_edge_density_check converts the region to grayscale, runs Canny edge detection, and computes the ratio of edge "
        "pixels to total pixels. This is high for thin, linear features like cracks.")
    add_paragraph(doc,
        "_directional_contrast_check computes Sobel gradients in X and Y. Cracks have a strong gradient in one direction and "
        "a weak gradient perpendicular to them. The function returns the anisotropy as one minus the ratio of smaller to larger "
        "mean gradient.")
    add_paragraph(doc,
        "_local_contrast_check builds a histogram of the region and compares the two most populated bins. A large difference "
        "indicates a bimodal distribution, which is typical for a dark crack on a light surface.")
    add_paragraph(doc,
        "_validate_crack_region combines the three checks with weights: edge 0.4, anisotropy 0.4, contrast 0.2. The minimum "
        "score is relaxed for small boxes because tiny cracks can have weaker CV signals.")

    add_heading(doc, "24.5 Non-Crack Object Detection", level=2)
    add_paragraph(doc,
        "_detect_non_crack_objects() runs the general YOLO model at a low confidence (0.25) and returns a list of boxes, "
        "confidences, and class names for objects that should not be counted as cracks. It also includes person, animal, and "
        "vehicle classes. These boxes are used during the primary, secondary, and fallback model passes to reject overlapping "
        "detections.")

    add_heading(doc, "24.6 The run_ensemble Function", level=2)
    add_paragraph(doc,
        "run_ensemble() is the main orchestrator. It takes an image array and a threshold. It first calls _detect_non_crack_objects. "
        "Then it runs the primary YOLO model. For each box it checks crack-like class, aspect ratio, area, overlap with non-crack "
        "objects, and CV validation. If the box passes, it extracts a polygon mask from the model's masks output and appends the "
        "detection with an adjusted confidence.")
    add_paragraph(doc,
        "The same logic applies to the secondary and fallback models. After collecting all candidates, the function combines them, "
        "sorts by confidence, and applies a custom NMS that is more permissive for small detections. The remaining boxes are "
        "merged with the non-crack objects and returned.")

    add_heading(doc, "24.7 Engineering Analysis", level=2)
    add_paragraph(doc,
        "Each detection goes through build_engineering_analysis(). It computes pixel-based dimensions using the pixel_scale_mm "
        "parameter, estimates a polygon area, maps the class to a Russian name, and assigns a severity. The overall_condition is "
        "set based on the most severe finding. The detailed response includes all of this metadata along with the original "
        "detection and an annotated image.")

    add_heading(doc, "24.8 PDF Report Generator", level=2)
    add_paragraph(doc,
        "report_generator.py uses ReportLab to create a PDF. The cover page includes the project name, inspector, date, and an "
        "annotated image. The summary page shows counts by severity and an overall condition. Subsequent pages list each finding "
        "with its dimensions, severity, and recommendations. The last page includes normative references and a disclaimer.")

    # Telegram bot deep dive
    add_heading(doc, "25. Telegram Bot in Detail", level=1)
    add_paragraph(doc,
        "The Telegram bot is a FastAPI application in the root main.py. It is separate from the web app and can be deployed on "
        "Vercel or any server. It uses a custom telegram_client.py to call the Telegram Bot API.")

    add_heading(doc, "25.1 Webhook Handler", level=2)
    add_paragraph(doc,
        "The /webhook/telegram endpoint reads the X-Telegram-Bot-Api-Secret-Token header and compares it to the configured secret. "
        "It parses the JSON body and extracts the message. It supports text, voice, and photo messages. Callback queries and other "
        "updates are acknowledged and ignored.")

    add_heading(doc, "25.2 Message Routing", level=2)
    add_paragraph(doc,
        "If the text is /start, it calls handle_start(). If a photo is present, it downloads the largest file and calls "
        "analyze_photo(). If a voice is present, it transcribes the voice and calls handle_message() with the transcript. For "
        "any other text it calls handle_message().")

    add_heading(doc, "25.3 Conversation State Machine", level=2)
    add_paragraph(doc,
        "handle_message() in bot/handlers.py is the main router. It loads the session for the chat_id and checks the state. If the "
        "state is HANDOFF, it waits for the 'bot' or 'бот' keyword to return. It also checks business hours and can hand off to a "
        "human if needed.")
    add_paragraph(doc,
        "The bot shows a menu with product options. When the user selects a product number, the bot either sends the corresponding "
        "PDF presentation or offers to connect to a manager. Product keys are mapped in _PRODUCT_PDF_MAP.")

    add_heading(doc, "25.4 InspectAI Photo Analysis", level=2)
    add_paragraph(doc,
        "bot/inspectai.py downloads the photo with httpx, posts it to /predict/detailed with environment and aggression parameters, "
        "and builds a Russian-language summary. It then posts the photo to /report with project metadata to generate a PDF and sends "
        "the PDF to the chat. Errors are caught and reported in Russian or Kazakh depending on the session language.")

    add_heading(doc, "25.5 AI Agent and CRM", level=2)
    add_paragraph(doc,
        "The AI agent (bot/ai_agent.py) sends messages to a Groq or Together model. It includes a system prompt that describes "
        "the company, products, and rules. The CRM module (bot/crm.py) posts lead data to Bitrix24 when the user asks to speak to "
        "a manager. The bot also uses language detection to respond in Russian or Kazakh.")

    # More practical chapters
    add_heading(doc, "26. Common Development Scenarios", level=1)

    add_heading(doc, "26.1 Adding a New Severity Level", level=2)
    add_paragraph(doc,
        "To add a new severity level, first update the Prisma enum Severity in schema.prisma and run db push. Then update the "
        "severityColors maps in dashboard pages, the SEV_COLORS and SEV_LABEL maps in the upload page, and the severity logic in "
        "the ML service. Finally update the PDF report generator to style the new severity.")

    add_heading(doc, "26.2 Adding a New Tool", level=2)
    add_paragraph(doc,
        "To add a new engineering tool, create a new static folder under ecosystem/, add the tool to the TOOL_URLS map in "
        "app/dashboard/tools/[tool]/page.tsx, add it to the tools array in app/dashboard/tools/page.tsx, and add it to the "
        "toolItems array in app/dashboard/layout.tsx.")

    add_heading(doc, "26.3 Connecting a Custom ML Model", level=2)
    add_paragraph(doc,
        "To use a custom YOLO model, place the .pt file in apps/ml-service/best.pt or models/best.pt, or upload it to a "
        "HuggingFace repository and set HF_MODEL and HF_MODEL_FILE. Ensure the model's class names are covered by "
        "_is_crack_like() and _is_non_crack_object().")

    add_heading(doc, "26.4 Migrating the Database", level=2)
    add_paragraph(doc,
        "Database changes follow the Prisma workflow: edit schema.prisma, run npx prisma generate to update the client, and "
        "npx prisma db push to apply changes to the database. For production, consider using npx prisma migrate deploy after "
        "creating migration files with npx prisma migrate dev.")

    add_heading(doc, "26.5 Customizing the Report", level=2)
    add_paragraph(doc,
        "The web PDF is built in lib/pdf-report.tsx using @react-pdf/renderer. You can add a logo, change colors, add sections, "
        "or include normative references. The ML PDF is built in apps/ml-service/report_generator.py using ReportLab. Update the "
        "styles, page templates, or content sections there.")

    add_heading(doc, "27. Testing and Quality Assurance", level=1)
    add_paragraph(doc,
        "InspectAI currently relies on manual testing. Recommended automated tests to add:")
    add_bullet(doc, "Unit tests for the ML service helpers: _is_crack_like, _validate_crack_region, NMS logic.")
    add_bullet(doc, "API route tests using Next.js with mock Prisma and session data.")
    add_bullet(doc, "End-to-end tests with Playwright for the upload, analysis, review, and report flows.")
    add_bullet(doc, "Performance tests for the ML service to measure inference time under load.")
    add_bullet(doc, "Accessibility tests to ensure color contrast and keyboard navigation.")

    add_heading(doc, "27.1 Manual Testing Checklist", level=2)
    add_numbered(doc, "Register a new account and log in.")
    add_numbered(doc, "Create a project with all optional fields.")
    add_numbered(doc, "Upload multiple photos to the project.")
    add_numbered(doc, "Run AI analysis with different thresholds and parameters.")
    add_numbered(doc, "Review each finding: confirm, reject, and edit with a note.")
    add_numbered(doc, "Generate a PDF report and verify the download link.")
    add_numbered(doc, "Open an engineering tool and test the fallback.")
    add_numbered(doc, "Test the public demo with a sample image.")
    add_numbered(doc, "Send a photo to the Telegram bot and verify the summary and PDF.")

    add_heading(doc, "28. Performance Considerations", level=1)
    add_paragraph(doc,
        "The ML service can be CPU-heavy. For production, deploy it on a GPU-enabled instance. The web app runs on Vercel's "
        "serverless functions, which have execution time limits. Long ML requests should be run against the external ML service, "
        "not inside Vercel, to avoid timeouts.")
    add_paragraph(doc,
        "Large images should be resized before upload. The current upload route does not resize images; adding a resize step would "
        "reduce ML inference time and storage costs. The dashboard page should also use pagination or virtualized lists if the "
        "number of analyses grows beyond a few hundred.")

    add_heading(doc, "29. Future Roadmap", level=1)
    add_paragraph(doc,
        "Potential improvements and extensions:")
    add_bullet(doc, "Email notifications when analyses complete (backend + email provider).")
    add_bullet(doc, "Team/organization accounts with multiple users per project.")
    add_bullet(doc, "Comparison of analyses over time for the same asset.")
    add_bullet(doc, "Mobile app for field inspections.")
    add_bullet(doc, "Integration with drones for automated photo capture.")
    add_bullet(doc, "Multi-language support for the web UI.")
    add_bullet(doc, "Model versioning and A/B testing.")
    add_bullet(doc, "Audit log for finding review status changes.")

    add_heading(doc, "30. Conclusion", level=1)
    add_paragraph(doc,
        "InspectAI is a full-stack AI product for structural defect detection. It combines a modern Next.js web app, a FastAPI "
        "machine-learning service, a conversational Telegram bot, and a set of engineering calculation tools. Each component has a "
        "clear responsibility and can be evolved independently.")
    add_paragraph(doc,
        "This handbook covered the architecture, technology stack, database schema, user interface, API routes, ML pipeline, "
        "Telegram bot, deployment, and common development scenarios. The goal was to provide enough detail for a new developer or "
        "engineer to understand every page, button, function, and service interaction. As the project grows, this document should "
        "be updated to reflect new features and changes.")


def add_extra_content(doc):
    """Append even more detailed content to reach 100+ pages."""

    # Part III
    add_heading(doc, "Part III: Code Walkthroughs and User Guides", level=1)
    add_paragraph(doc,
        "This part contains line-by-line explanations of important files and step-by-step guides for the end user. "
        "The purpose is to bridge the gap between high-level architecture and implementation details, so both developers "
        "and operators can learn from the same document.")

    add_heading(doc, "31. Dashboard Layout Walkthrough", level=1)
    add_paragraph(doc,
        "The dashboard layout is one of the most important client files. It is a client component because it uses "
        "NextAuth session data and React state for the mobile drawer. The file begins with a list of navigation "
        "items and tool items. These arrays are the single source of truth for the left sidebar.")
    add_paragraph(doc,
        "The navItems array contains objects with href, icon, and label. The icon value is a Material Symbols name. "
        "These are rendered as text inside a span with the material-symbols-outlined class. The toolItems array is "
        "rendered under the 'Engineering Tools' heading. The isTool flag is passed to NavLink so that nested routes "
        "are correctly marked as active.")
    add_paragraph(doc,
        "NavLink is a pure rendering component. It uses the active boolean to decide between the active and inactive "
        "styles. The active computation for tools uses pathname.startsWith(item.href), while main items use an exact "
        "match or a startsWith check for non-dashboard routes. This ensures /dashboard/projects/[id] keeps 'Projects' "
        "highlighted.")
    add_paragraph(doc,
        "The Logo function is a small reusable component. It returns a Link to /dashboard with the InspectAI letter I "
        "inside a colored box. This is rendered at the top of both desktop and mobile sidebars.")
    add_paragraph(doc,
        "DashboardLayout uses three useState hooks: mounted, mobileOpen, and relies on useSession. The first useEffect "
        "sets mounted to true on the client. The second useEffect redirects to /login if the session is unauthenticated "
        "after the component has mounted. The third useEffect closes the mobile drawer when the pathname changes.")
    add_paragraph(doc,
        "If the session is loading or not present, a full-screen spinner is shown. This prevents the dashboard content "
        "from rendering before the user is known. Once loaded, the desktop sidebar and the main content area are rendered. "
        "The main content has a sticky top bar and a padded container for children.")
    add_code_block(doc,
"""const navItems = [
  { href: "/dashboard", icon: "dashboard", label: "Dashboard" },
  { href: "/dashboard/projects", icon: "folder", label: "Projects" },
  { href: "/dashboard/upload", icon: "cloud_upload", label: "Upload Assets" },
  { href: "/dashboard/analysis", icon: "psychology", label: "Analysis Engine" },
  { href: "/dashboard/review", icon: "fact_check", label: "Review Queue" },
  { href: "/dashboard/reports", icon: "description", label: "Final Reports" },
  { href: "/dashboard/tools", icon: "construction", label: "Engineering Tools" },
  { href: "/dashboard/knowledge", icon: "menu_book", label: "Knowledge Base" },
];""")

    add_heading(doc, "32. Dashboard Home Walkthrough", level=1)
    add_paragraph(doc,
        "The dashboard home uses a useEffect to load data from /api/dashboard. It defines a DashboardData interface "
        "that matches the JSON returned by the API. The kpis array combines the API values with trend strings and "
        "visual bar widths.")
    add_paragraph(doc,
        "Each KPI is a card with a label in label-caps style, a large numeric value, a trend badge, and a progress bar. "
        "The PENDING REVIEWS card has a conditional 'HIGH ATTENTION' banner when the value is greater than zero. This "
        "is a visual cue to encourage users to review findings.")
    add_paragraph(doc,
        "The severity distribution section uses an IIFE (immediately invoked function expression) inside the JSX to "
        "compute the conic gradient. This is a compact pattern, but it can be extracted into a separate component in "
        "a future refactor. The gradient is built by iterating over the severity counts, converting them to angles, "
        "and joining them into a conic-gradient CSS string.")
    add_paragraph(doc,
        "The activity timeline uses another IIFE to compute bar heights relative to the maximum count. The days are "
        "derived from the date strings using getUTCDay(). The chart is not an SVG or canvas chart; it is built from "
        "plain divs, which keeps the bundle small and the rendering fast.")
    add_paragraph(doc,
        "Recent projects are rendered in a table. Each row is a Link to the project detail. The objectType, photo count, "
        "last activity date, and status are shown. Status is colored with Tailwind background and text classes. If there "
        "are no projects, a CTA row links to the project creation page.")

    add_heading(doc, "33. Upload Page Walkthrough", level=1)
    add_paragraph(doc,
        "The upload page has the most complex state. It tracks projects, selected project, files, drag state, loading, "
        "processing, progress, result message, selected detail, active tool, and analysis parameters. The interface "
        "definitions at the top of the file help TypeScript check the data shapes.")
    add_paragraph(doc,
        "handleFiles is wrapped with useCallback. It receives a FileList from either the drag-and-drop event or the "
        "file input. It maps each file into an UploadedFile object with a preview URL. The preview is rendered as a "
        "thumbnail for the user.")
    add_paragraph(doc,
        "The uploadAll function is the core processing logic. It loops over the files array sequentially. For each file "
        "it updates the status to 'uploading', posts the file to /api/assets/upload, receives the asset, and then posts "
        "to /api/analyses/run with the assetId, projectId, and analysisParams. If upload fails, the status becomes 'error'. "
        "If analysis fails, it also becomes 'error'. On success it becomes 'done' and stores the summary.")
    add_paragraph(doc,
        "The progress bar is updated at the start of each iteration using Math.round((i / total) * 100). This gives the "
        "user a rough sense of completion. The resultMessage shows a short summary after all files are processed.")
    add_paragraph(doc,
        "The detail view shows the selected file's image with finding overlays. It reuses the same overlay logic as the "
        "project detail and analysis detail pages: the image is scaled to the container, and each finding is rendered "
        "with a positioned div or SVG polygon. The user can click a finding to open the CrackCalc tool pre-filled with "
        "the finding's dimensions.")

    add_heading(doc, "34. Analysis Run API Walkthrough", level=1)
    add_paragraph(doc,
        "The /api/analyses/run route is the bridge between the web app and the ML service. It is a server-side Next.js "
        "route handler. The first action is to verify the session with getServerSession. If no session, it returns 401.")
    add_paragraph(doc,
        "After validating the input body, it finds the asset and includes its project relation. The ownership check "
        "compares asset.project.userId with the session user id. If they do not match, it returns 403. This prevents "
        "one user from analyzing another user's assets.")
    add_paragraph(doc,
        "It then creates the Analysis record with status PROCESSING. This is important because the ML call can take "
        "several seconds. The user sees a PROCESSING status in the UI. Next it fetches the image bytes from the Vercel "
        "Blob URL. The fetch is straightforward: fetch(asset.blobUrl) and then blob().")
    add_paragraph(doc,
        "The FormData is constructed with the image blob and filename. The search parameters are built from the whitelist. "
        "For each whitelisted key, if the params object contains a non-empty value, it is appended. This prevents undefined "
        "or empty strings from being sent. The ML endpoint becomes something like:")
    add_code_block(doc,
"""https://alllxndr-inspectai-ml.hf.space/predict/detailed?
  pixel_scale_mm=0.05&environment=atmospheric&aggression=normal&threshold=0.25""")
    add_paragraph(doc,
        "The ML response is parsed. The annotated_image field is stripped because it is a large base64 string. The "
        "detections are filtered to exclude the 'other' class. For each remaining detection, a Finding row is created. "
        "The severity is uppercased to match the enum. The engineering dimensions come from result.engineering.")
    add_paragraph(doc,
        "Finally, the Analysis is updated to COMPLETED with confidence, modelVersion, and the lean resultData. If an "
        "error occurs at any point, the Analysis is updated to FAILED and a 500 response with the error message is returned.")

    add_heading(doc, "35. User Step-by-Step Guide", level=1)
    add_paragraph(doc,
        "This chapter is written for engineers who will use the platform in the field. It explains how to complete a "
        "typical inspection workflow from registration to PDF report.")
    add_numbered(doc, "Open the web app and click 'Sign in' or 'Create one' to register.")
    add_numbered(doc, "After registration, you are taken to the Dashboard. Click 'CREATE PROJECT' on the Projects page.")
    add_numbered(doc, "Enter the project name, optional site ID, object type, address, and description. Click CREATE.")
    add_numbered(doc, "Navigate to Upload Assets. Select the project from the dropdown.")
    add_numbered(doc, "Optionally expand the ANALYSIS PARAMETERS section and set the pixel scale and threshold.")
    add_numbered(doc, "Drag and drop inspection photos or click the drop zone to select files.")
    add_numbered(doc, "Click UPLOAD & ANALYZE. Wait for the progress bar to reach 100%.")
    add_numbered(doc, "Click a completed photo to view the findings overlay.")
    add_numbered(doc, "Go to Review Queue and review each PENDING finding. Confirm correct detections, reject false positives, or edit with a note.")
    add_numbered(doc, "Go to Final Reports and click GENERATE REPORT. Select the project. Download the PDF.")
    add_numbered(doc, "If needed, open CrackCalc or other engineering tools from the sidebar and enter dimensions from the findings.")

    add_heading(doc, "36. ML Algorithm Deep Dive", level=1)
    add_paragraph(doc,
        "This chapter explains how the ML service detects structural defects. The approach is a hybrid of deep learning "
        "and classical computer vision. The deep learning model proposes regions of interest. The classical vision step "
        "validates whether each region actually looks like a crack.")
    add_paragraph(doc,
        "YOLOv8 is an anchor-free object detector. It divides the image into a grid and predicts bounding boxes and class "
        "probabilities directly. The primary YOLO model is trained on concrete defect images. The secondary and fallback "
        "models provide additional coverage for different crack shapes and surface types.")
    add_paragraph(doc,
        "The general YOLO model is a pretrained COCO model. It is not used to find cracks. Instead, it is used to find "
        "objects that are commonly confused with cracks: pipes, wires, windows, shadows, stains, and so on. Any proposed "
        "crack that overlaps with one of these objects is discarded.")
    add_paragraph(doc,
        "After the deep learning proposal, the region is validated with OpenCV. Canny edge detection highlights edges. "
        "Sobel gradients measure directionality. A histogram check measures contrast. The combined score must exceed a "
        "threshold. The threshold is lower for small boxes because tiny cracks are harder to see.")
    add_paragraph(doc,
        "Non-Maximum Suppression removes overlapping boxes. The custom NMS is size-aware: small detections are allowed "
        "to overlap more than large ones because thin cracks can run close together. The final detections are returned "
        "with bounding boxes and optional polygon masks.")
    add_paragraph(doc,
        "Engineering analysis converts pixel dimensions to millimeters using the pixel_scale_mm parameter. It also "
        "assigns a severity and an overall condition. The result is a JSON object suitable for storage in the database "
        "and for rendering in the UI.")

    add_heading(doc, "37. Telegram Bot User Guide", level=1)
    add_paragraph(doc,
        "Users can also interact with InspectAI through Telegram. This is useful for quick field inspections when the "
        "engineer does not want to open a web browser.")
    add_numbered(doc, "Open the Telegram bot and start the chat with /start.")
    add_numbered(doc, "Select a city and choose a menu option if prompted.")
    add_numbered(doc, "Send a photo of the concrete surface.")
    add_numbered(doc, "Wait a few seconds. The bot will return a Russian-language summary with counts and severities.")
    add_numbered(doc, "A few moments later the bot will send a PDF inspection report.")
    add_numbered(doc, "If you have questions, send a text message. If the bot cannot answer, it can hand off to a manager.")

    add_heading(doc, "38. Common Customizations", level=1)
    add_paragraph(doc,
        "This chapter describes frequent customization requests and where to make changes.")
    add_bullet(doc, "Change the landing page copy: edit the static arrays in app/web/app/page.tsx.")
    add_bullet(doc, "Change colors: edit the colors object in tailwind.config.ts.")
    add_bullet(doc, "Add a new navigation item: add it to navItems and optionally toolItems in app/dashboard/layout.tsx.")
    add_bullet(doc, "Add a new API endpoint: create route.ts under app/api/<path> and import authOptions and prisma.")
    add_bullet(doc, "Add a new defect class: update _is_crack_like in apps/ml-service/main.py and add styling in the web UI.")
    add_bullet(doc, "Change the report language: edit lib/pdf-report.tsx for web reports and report_generator.py for ML reports.")
    add_bullet(doc, "Connect a different LLM provider: edit bot/config.py and bot/ai_agent.py.")

    add_heading(doc, "39. Monitoring and Logging", level=1)
    add_paragraph(doc,
        "InspectAI uses basic logging. The Python services log to stdout. Vercel captures logs from Next.js API routes. "
        "For production, consider adding a log aggregation service such as Datadog, Sentry, or a simple cloud logging "
        "solution.")
    add_paragraph(doc,
        "Important metrics to monitor:")
    add_bullet(doc, "ML service response time — should be under a few seconds for typical images.")
    add_bullet(doc, "API error rate — especially for /api/analyses/run and /api/reports/generate.")
    add_bullet(doc, "Database connection health — Prisma can log slow queries.")
    add_bullet(doc, "Blob storage usage — Vercel Blob has limits and costs.")
    add_bullet(doc, "Telegram webhook delivery — Telegram retries failed deliveries.")

    add_heading(doc, "40. Backup and Data Retention", level=1)
    add_paragraph(doc,
        "PostgreSQL backups should be configured on the hosting provider. Neon, AWS RDS, and similar services offer "
        "automated backups. Vercel Blob objects are not automatically backed up; consider a periodic export if reports "
        "are critical.")
    add_paragraph(doc,
        "Data retention policy should be defined by the organization. InspectAI stores user profiles, projects, assets, "
        "analyses, findings, and reports. Old analyses can be archived or deleted based on project status. The Report "
        "table should be the source of truth for generated PDFs; deleting the blob URL will make the download link fail, "
        "but the report metadata remains.")


def add_code_reference(doc):
    """Read and include key source files for completeness."""

    def include_file(rel_path, heading_text, description=""):
        full_path = f"/Users/aleksandr/engeenering-ml-fasad/{rel_path}"
        try:
            with open(full_path, "r", encoding="utf-8") as f:
                code = f.read()
        except Exception as e:
            code = f"[Could not read {rel_path}: {e}]"
        add_heading(doc, heading_text, level=2)
        if description:
            add_paragraph(doc, description)
        add_paragraph(doc, f"File: {rel_path}")
        p = doc.add_paragraph()
        p.style = "Normal"
        run = p.add_run(code[:15000] if len(code) > 15000 else code)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        if len(code) > 15000:
            add_paragraph(doc, "(File truncated for brevity; see the repository for the full source.)")

    add_heading(doc, "Part IV: Key Source Files", level=1)
    add_paragraph(doc,
        "This part reproduces the most important source files for reference. Reading the actual code alongside the "
        "explanations in the previous chapters helps connect documentation to implementation. Some long files are truncated.")

    include_file(
        "apps/web/lib/auth.ts",
        "41. Authentication Configuration (lib/auth.ts)",
        "The NextAuth options define the credentials provider, session strategy, and callbacks that attach the user id and role to the JWT.")

    include_file(
        "apps/web/prisma/schema.prisma",
        "42. Database Schema (prisma/schema.prisma)",
        "The Prisma schema defines all database tables, enums, and relationships used by the web application.")

    include_file(
        "apps/web/app/dashboard/layout.tsx",
        "43. Dashboard Layout (app/dashboard/layout.tsx)",
        "The full dashboard layout component with navigation, mobile drawer, and top bar.")

    include_file(
        "apps/web/app/dashboard/page.tsx",
        "44. Dashboard Home (app/dashboard/page.tsx)",
        "The dashboard home page with KPIs, severity distribution, activity timeline, and recent projects.")

    include_file(
        "apps/web/app/api/analyses/run/route.ts",
        "45. Analysis Run API (app/api/analyses/run/route.ts)",
        "The API route that fetches the asset image, sends it to the ML service, and stores findings.")

    include_file(
        "apps/web/app/api/reports/generate/route.ts",
        "46. Report Generation API (app/api/reports/generate/route.ts)",
        "The API route that aggregates project data, renders a PDF, and uploads it to Vercel Blob.")

    include_file(
        "apps/web/app/api/dashboard/route.ts",
        "47. Dashboard Summary API (app/api/dashboard/route.ts)",
        "The API route that computes dashboard statistics, severity distribution, and weekly activity.")

    include_file(
        "apps/ml-service/main.py",
        "48. ML Service (apps/ml-service/main.py)",
        "The FastAPI ML service. This file is long and is included in full as a reference.")

    include_file(
        "bot/inspectai.py",
        "49. Telegram InspectAI Integration (bot/inspectai.py)",
        "Handles photo download, ML inference, and PDF generation for the Telegram bot.")

    include_file(
        "main.py",
        "50. Telegram Bot Main Entry (main.py)",
        "The FastAPI entry point for the Telegram bot, including the webhook endpoint.")

    include_file(
        "bot/config.py",
        "51. Bot Configuration (bot/config.py)",
        "Pydantic Settings configuration showing all environment variables for the bot.")

    include_file(
        "apps/web/app/dashboard/projects/[id]/page.tsx",
        "52. Project Detail Page (app/dashboard/projects/[id]/page.tsx)",
        "Full project detail page with asset gallery, analysis parameters, and report generation.")

    include_file(
        "apps/web/app/dashboard/projects/[id]/AssetModal.tsx",
        "53. Asset Modal (app/dashboard/projects/[id]/AssetModal.tsx)",
        "Modal for viewing an asset image with finding overlays and review controls.")

    include_file(
        "apps/web/app/dashboard/upload/page.tsx",
        "54. Batch Upload Page (app/dashboard/upload/page.tsx)",
        "Full batch upload and analysis page. Long file truncated.")

    include_file(
        "apps/web/app/dashboard/review/page.tsx",
        "55. Review Queue (app/dashboard/review/page.tsx)",
        "The human-in-the-loop review queue for confirming, rejecting, and editing findings.")

    include_file(
        "apps/web/app/dashboard/settings/page.tsx",
        "56. Settings Page (app/dashboard/settings/page.tsx)",
        "User profile, ML threshold, and notification settings.")

    include_file(
        "apps/web/app/demo/page.tsx",
        "57. Public Demo Page (app/demo/page.tsx)",
        "Public single-image analysis and PDF download demo.")

    include_file(
        "apps/web/app/page.tsx",
        "58. Landing Page (app/page.tsx)",
        "Marketing landing page with hero, features, FAQ, and CTAs.")

    include_file(
        "apps/web/app/api/findings/[id]/route.ts",
        "59. Finding Update API (app/api/findings/[id]/route.ts)",
        "PATCH endpoint with ownership check for updating finding review status.")

    include_file(
        "apps/web/app/api/projects/route.ts",
        "60. Projects API (app/api/projects/route.ts)",
        "GET and POST routes for listing and creating projects.")

    include_file(
        "apps/ml-service/report_generator.py",
        "61. ReportLab Report Generator (apps/ml-service/report_generator.py)",
        "Generates the ML service PDF report. Very long file truncated.")


if __name__ == "__main__":
    main()
